import streamlit as st
import io
from docx import Document
from orchestrator import generate

def create_docx(contract_data):
    doc = Document()
    doc.add_heading('Freelance Contract', 0)
    
    sections = [
        ("Parties", contract_data.get("Parties", "")),
        ("Scope of Work", contract_data.get("Scope", "")),
        ("Deliverables", contract_data.get("Deliverables", "")),
        ("Payment Terms", contract_data.get("Payment", "")),
        ("Revisions", contract_data.get("Revisions", "")),
        ("Intellectual Property", contract_data.get("IP", "")),
        ("Dispute Resolution", contract_data.get("Dispute Resolution", ""))
    ]
    
    for title, content in sections:
        if content:
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def render(user):
    st.header("Contract Drafter")
    st.write("Generate a professional freelance contract. Export as .docx.")
    
    with st.form("contract_form"):
        title = st.text_input("Project Title")
        desc = st.text_area("Project Description")
        
        col1, col2 = st.columns(2)
        with col1:
            payment_terms = st.text_input("Payment Terms (e.g., $500 total, 50% upfront)")
            milestones = st.selectbox("Structure", ["Full Project", "Milestone Based"])
        with col2:
            revisions = st.number_input("Number of Revisions", min_value=0, value=2)
            client_name = st.text_input("Client Name (Optional)")
            
        submit = st.form_submit_button("Generate Contract")
        
    if submit:
        if not title or not desc:
            st.warning("Please provide a Project Title and Description.")
            return
            
        with st.spinner("Drafting contract..."):
            prompt = f"""
            Draft a professional freelance contract.
            Project Title: {title}
            Description: {desc}
            Client Name: {client_name or '[Client Name]'}
            Freelancer Name: {user.get('name')}
            Payment Terms: {payment_terms}
            Structure: {milestones}
            Revisions Included: {revisions}
            
            Return a JSON object with exactly these keys containing the text for each section:
            - "Parties"
            - "Scope"
            - "Deliverables"
            - "Payment"
            - "Revisions"
            - "IP"
            - "Dispute Resolution"
            """
            
            result = generate(user, prompt)
            
            if "error" in result:
                st.error("Error drafting contract:")
                st.write(result["error"])
                if "raw_response" in result:
                    st.text_area("Raw Response", result["raw_response"])
            else:
                st.success("Contract Generated Successfully!")
                
                # Preview
                with st.expander("Preview Contract", expanded=True):
                    st.markdown(f"### Parties\n{result.get('Parties', '')}")
                    st.markdown(f"### Scope of Work\n{result.get('Scope', '')}")
                    st.markdown(f"### Deliverables\n{result.get('Deliverables', '')}")
                    st.markdown(f"### Payment Terms\n{result.get('Payment', '')}")
                    st.markdown(f"### Revisions\n{result.get('Revisions', '')}")
                    st.markdown(f"### Intellectual Property\n{result.get('IP', '')}")
                    st.markdown(f"### Dispute Resolution\n{result.get('Dispute Resolution', '')}")
                
                # Export to DOCX
                docx_data = create_docx(result)
                st.download_button(
                    label="Download Contract (.docx)",
                    data=docx_data,
                    file_name=f"Contract_{title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
